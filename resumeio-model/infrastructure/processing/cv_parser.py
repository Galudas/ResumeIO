import os
import random
import docx
import regex
from infrastructure.aggregate import applicant

cv_file_path = "../resumeio-generator/cv_data/"
position_file_path = "../resumeio-generator/positions.txt"


def random_position(file_name):
    file = open(file_name, "r")
    line = next(file)
    for num, aline in enumerate(file, 2):
        if random.randrange(num):
            continue
        line = aline
    return line


def read_docx():
    position_name = random_position(position_file_path)
    file_regex = "cv_" + position_name.replace("\n", "") + "_.*"
    cv_list = [file for file in os.listdir(cv_file_path) if regex.match(file_regex, file)]
    for cv in cv_list:
        print(cv_data(cv_file_path + cv))


def cv_data(filename):
    doc = docx.Document(filename)
    last_name = ""
    firs_name = ""
    email = ""
    position = ""
    experience = ""
    skills = ""
    i = 0
    while i < len(doc.paragraphs):
        if doc.paragraphs[i].text.find("Last name:") != -1:
            last_name = doc.paragraphs[i].text.split("Last name: ")[1]
        if doc.paragraphs[i].text.find("First name:") != -1:
            firs_name = doc.paragraphs[i].text.split("First name: ")[1]
        if doc.paragraphs[i].text.find("Email:") != -1:
            email = doc.paragraphs[i].text.split("Email: ")[1]
        if doc.paragraphs[i].text.find("Experience:") != -1:
            position = doc.paragraphs[i].text.split("Experience: ")[1]
            i += 1
            while doc.paragraphs[i].text != 'Skills:':
                experience += doc.paragraphs[i].text
                i += 1
        if doc.paragraphs[i].text.find("Skills:") != -1:
            i += 1
            while i < len(doc.paragraphs):
                skills += doc.paragraphs[i].text
                i += 1
        i += 1
    return applicant.Applicant(experience, skills, firs_name + " " + last_name, email, position)


read_docx()