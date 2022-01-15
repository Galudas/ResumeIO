import os
import random
import docx
from infrastructure.aggregate import job

job_description_file_path = '../resumeio-generator/job_data'


def get_random_job():
    nr = random.randint(1, len(os.listdir(job_description_file_path)))
    return job_description_file_path + '/' + 'job_' + str(nr) + '.docx'


def parse_job():
    file = get_random_job()
    print(file)
    doc = docx.Document(file)
    department = ''
    description = ''
    requirements = ''
    i = 0
    while i < len(doc.paragraphs):
        if doc.paragraphs[i].text.find("Department:") != -1:
            department = doc.paragraphs[i].text.split("Department:")[1]
            if department == '':
                department = doc.paragraphs[0]
        if doc.paragraphs[i].text.find("Description:") != -1:
            i += 1
            while doc.paragraphs[i].text != "Requirements:":
                description += doc.paragraphs[i].text
                i += 1
        if doc.paragraphs[i].text.find("Requirements:") != -1:
            i += 1
            while i < len(doc.paragraphs):
                requirements += doc.paragraphs[i].text
                i += 1
        i += 1
    return job.Job(description, department, requirements)


print(parse_job())
