from docx import Document
import csv
import names


def generate_cvs():
    with open('linkedin.csv', encoding='utf-8') as csvfile:
        position_list = []
        csv_file = csv.reader(csvfile)
        line_count = 0
        for row in csv_file:
            if line_count != 0:
                name = names.get_full_name()
                document = Document()

                document.add_heading('CV', 0)

                document.add_heading('Personal data', level=1)
                document.add_paragraph('Last name: ' + name.split(" ")[0])
                document.add_paragraph('First name: ' + name.split(" ")[1])
                document.add_paragraph('Email: ' + name.split(" ")[0].lower() + "_" + name.split(" ")[1].lower() + "@"
                                       + "domain.com")
                if row[1] not in position_list:
                    position_list.append(row[1])
                document.add_heading('Experience: ' + row[1], level=1)
                document.add_paragraph(row[5])

                document.add_heading('Skills:', level=1)
                document.add_paragraph(row[9].replace("[", "").replace("]", "").replace("'\\n", "").replace("\\n'", ""))
                document.save(f'cv_data/cv_{row[1] + "_" + str(line_count)}.docx')
            line_count += 1
        f = open("positions.txt", "w+")
        print(position_list)
        for position in position_list:
            f.write(str(position) + "\n")
        f.close()


generate_cvs()
