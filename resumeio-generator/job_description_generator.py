from docx import Document
import csv


def generate_job_description():
    with open('fake_job_postings.csv', encoding='utf-8') as csvfile:
        job_list = []
        csv_file = csv.reader(csvfile)
        line_count = 0
        for row in csv_file:
            if line_count != 0:
                if row[1] not in job_list:
                    job_list.append(row[1])
                document = Document()
                document.add_heading(row[1], 0)
                document.add_heading('Department: ' + row[3], level=1)
                document.add_heading('Description:', level=1)
                document.add_paragraph(row[6])
                document.add_heading('Requirements:', level=1)
                document.add_paragraph(row[7])
                document.save(f'job_data/job_{str(line_count)}.docx')
            line_count += 1


generate_job_description()
